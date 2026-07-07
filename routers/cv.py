from fastapi import APIRouter, UploadFile, File, Form, HTTPException
from models.schemas import CVAnalysisResponse
from services.pdf_parser import extract_text_from_pdf
from services.analyze import extract_skills, analyze_career_paths, analyze_cv_ats, detect_experience_level
from services.gemini_service import analyze_cv_with_gemini
import shutil
import os
from langdetect import detect
from docx import Document

router = APIRouter()

UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

# 🔥 دالة قراءة ملفات الـ Word بسرعة وبدون OCR
def extract_text_from_docx(file_path):
    doc = Document(file_path)
    valid_blocks = []
    
    # 1. استخراج النصوص من الفقرات العادية
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            valid_blocks.append([None, text])
    
    # 2. 🔥 الحل: استخراج النصوص من داخل الجداول
    # هذا سيضمن قراءة أي مهارات موجودة في جداول
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # قراءة نص كل خلية وتنظيفه
                cell_text = cell.text.strip()
                if cell_text:
                    valid_blocks.append([None, cell_text])
                    
    return valid_blocks

@router.post("/analyze", response_model=CVAnalysisResponse)
async def analyze_cv(
    file: UploadFile = File(...),
    career_path: str = Form("backend"),
  #  career_path: str = Form(...)
):
    
    print(f"Received file: {file.filename}, Career Path: {career_path}")

    mapping = {
        "frontend": "Frontend Developer",
        "backend": "Backend Developer",
        "fullstack": "Full Stack Developer",
        "network": "Network Engineer",
        "ai": "AI Engineer"
    }

    if not file.filename.endswith(".pdf") and not file.filename.endswith(".docx"):
        raise HTTPException(status_code=400, detail="الملف يجب أن يكون بصيغة PDF أو DOCX")

    file_path = os.path.join(UPLOAD_DIR, file.filename)
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)


    try:
        # 1. استخراج النصوص
        if file.filename.endswith(".docx"):
            ocr_data = extract_text_from_docx(file_path)
        else:
            ocr_data = extract_text_from_pdf(file_path)

        if not ocr_data:
            raise HTTPException(status_code=422, detail="تعذر استخراج أي نص من الملف. يرجى التأكد من أن الملف ليس صورة فارغة أو PDF محمي.")

        # 🔥 تجهيز النص ككتلة واحدة (String) لإرساله لدالة الـ ATS
        full_text = " ".join([item[1] for item in ocr_data if len(item) > 1 and item[1]])

        try:
             detected_language = detect(full_text)
        except:
            detected_language = "unknown"


        # 2. استخراج المهارات
        skills = extract_skills(ocr_data)

        # 🔥 استنتاج مستوى الخبرة (junior/senior) من النص والمهارات المستخرجة
        experience_level = detect_experience_level(full_text, skills)

        # 3. استدعاء الخوارزمية (بناءً على متطلبات المستوى المستنتج)
        all_analysis = analyze_career_paths(skills, experience_level=experience_level)
        
        target_title = mapping.get(career_path.lower(), "Frontend Developer")

        # 4. البحث عن النتيجة 
        selected_path_data = None
        if all_analysis and isinstance(all_analysis, dict) and "career_paths" in all_analysis:
            for path in all_analysis["career_paths"]:
                if str(path.get("title", "")).lower() == target_title.lower():
                    selected_path_data = path
                    break
        
        if not selected_path_data:
            raise HTTPException(status_code=404, detail="لم يتم العثور على مهارات متوافقة مع هذا المسار الوظيفي.")

        # 🔥 5. استدعاء خوارزمية الـ ATS بناءً على النص والمهارات
        # تجميع كل المهارات المطلوبة للمسار (سواء مطابقة أو ناقصة) لتمريرها للدالة
        all_path_skills = selected_path_data.get("matched_skills", []) + selected_path_data.get("missing_skills", [])
        # استخراج اسم المهارة فقط في حال كانت Object
        expected_skill_names = [s if isinstance(s, str) else s.get("skill", "") for s in all_path_skills]

        ats_report = analyze_cv_ats(
            pdf_path=file_path,
            extracted_text=full_text,
            expected_skills=expected_skill_names,
            ocr_data=ocr_data
        )

        ai_assessment = analyze_cv_with_gemini(
            extracted_text=full_text,
            extracted_skills=skills,
            chosen_path=mapping.get(career_path, career_path)
        )

        # استخراج بيانات الـ ATS بشكل آمن تماماً سواء كان ats_report كائن (Object) أو قاموس (dict)
        if isinstance(ats_report, dict):
            ats_dict = {
                "is_compliant": ats_report.get("is_compliant", False),
                "score": ats_report.get("score", 0),
                "issues": ats_report.get("issues", []),
                "passed": ats_report.get("passed", []),
                "ai_ats_report": ai_assessment.ai_ats_report
            }
        else:
            ats_dict = {
                "is_compliant": getattr(ats_report, "is_compliant", False),
                "score": getattr(ats_report, "score", 0),
                "issues": getattr(ats_report, "issues", []),
                "passed": getattr(ats_report, "passed", []),
                "ai_ats_report": ai_assessment.ai_ats_report
            }

        # 6. الرد الصافي للفرونت إند (مع تقرير الـ ATS)
        return CVAnalysisResponse(
            detected_language=detected_language,
            extracted_skills=skills,
            experience_level=ai_assessment.experience_level,
            career_paths=[{
                "title": selected_path_data.get("title", target_title),
                "match_score": int(selected_path_data.get("match_score", 0)),
                "matched_skills": selected_path_data.get("matched_skills", []),
                "missing_skills": selected_path_data.get("missing_skills", [])
            }],
            top_recommendation=all_analysis.get("top_recommendation", target_title),
            is_path_suitable=ai_assessment.is_path_suitable,   # حقل جديد من Gemini
            ai_path_feedback=ai_assessment.ai_path_feedback,
            ats_analysis=ats_dict
        )

    except Exception as e:
        print(f"❌ خطأ بالباكيند: {str(e)}")
        raise HTTPException(status_code=500, detail=f"خطأ في معالجة الملف: {str(e)}")

    finally:
        if os.path.exists(file_path):
            os.remove(file_path)